from __future__ import annotations

from pathlib import Path
import zipfile

from docx import Document
from main import build_content_disposition

from text_slicer import slice_text
from rules.punct_rules import check_punctuation_rules
from rules.format_rules import check_format_rules
from docx_utils import build_final_docx_from_text
from pipeline import analyze
import pipeline


def test_slice_text_splits_and_preserves_order() -> None:
    text = "第一段" + "a" * 30 + "\n" + "第二段" + "b" * 30 + "\n" + "第三段" + "c" * 30
    chunks = slice_text(text, max_chars=50)
    assert len(chunks) >= 2
    assert "第一段" in chunks[0]
    assert "第三段" in "\n".join(chunks)


def test_punctuation_rule_detects_multiple_repeat_punct_in_one_paragraph() -> None:
    paragraphs = ["这个句子。。还有这个！！再来一个？？"]
    issues = check_punctuation_rules(paragraphs)
    repeat_issues = [i for i in issues if "连续重复符号" in i.message]
    assert len(repeat_issues) >= 3


def test_format_rules_warn_on_empty_content(tmp_path: Path) -> None:
    empty_docx = tmp_path / "empty.docx"
    Document().save(empty_docx)
    issues = check_format_rules(str(empty_docx), [])
    assert any(i.issue_type == "结构" and "文档内容为空" in i.message for i in issues)


def test_analyze_text_mode_produces_expected_files() -> None:
    result = analyze(mode="text", text="关于测试通知\n各部门：\n请按时完成。\n2026年2月14日")
    assert "标记后文档.docx" in result["files"]
    assert isinstance(result["annotated_bytes"], (bytes, bytearray))


def test_analyze_user_docx_file_produces_outputs() -> None:
    input_path = Path(__file__).resolve().parents[2] / "锦江中心年中工作会材料0805会后修改.docx"
    assert input_path.exists(), f"missing file: {input_path}"
    result = analyze(mode="file", file_path=str(input_path), do_polish=False)
    assert isinstance(result["report"], dict)
    assert "issues" in result["report"]
    assert "标记后文档.docx" in result["files"]


def test_build_final_docx_from_text_creates_file(tmp_path: Path) -> None:
    out = tmp_path / "定稿正文.docx"
    build_final_docx_from_text("标题\n正文第一段", str(out))
    assert out.exists() and out.stat().st_size > 0


def test_quote_chain_rule_detects_punct_between_quotes() -> None:
    paragraphs = ["“内容1”、“内容2”"]
    issues = check_punctuation_rules(paragraphs)
    assert any(i.issue_type == "引号" and "并列引号内容之间不应加标点" in i.message for i in issues)


def test_format_rules_detect_end_punct_and_space(tmp_path: Path) -> None:
    from docx_utils import build_final_docx_from_text
    out = tmp_path / "x.docx"
    build_final_docx_from_text("标题\n这是 有 空格的段落", str(out))
    issues = check_format_rules(str(out), ["标题", "这是 有 空格的段落"])
    assert any(i.issue_type == "空格" for i in issues)


def test_annotated_doc_keeps_issue_paragraph_alignment_with_blank_lines(tmp_path: Path) -> None:
    from docx_utils import build_annotated_original_docx
    from models import Issue

    src = tmp_path / "src.docx"
    doc = Document()
    doc.add_paragraph("第一段")
    doc.add_paragraph("")
    doc.add_paragraph("第二段。。")
    doc.save(src)

    out = tmp_path / "out.docx"
    issues = [Issue(issue_type="标点", severity="error", paragraph_index=2, message="疑似存在连续重复符号：。。。", evidence="。。", suggestion="删除多余符号")]
    build_annotated_original_docx(str(src), issues, str(out))

    new_doc = Document(out)
    assert "【疑似存在连续重复符号：。。。】" in new_doc.paragraphs[2].text
    assert "第一段" == new_doc.paragraphs[0].text


def test_build_content_disposition_uses_rfc5987_encoding() -> None:
    header = build_content_disposition("标记后文档.docx")
    assert "filename=annotated.docx" in header
    assert "filename*=UTF-8''" in header
    assert "%E6%A0%87%E8%AE%B0%E5%90%8E%E6%96%87%E6%A1%A3.docx" in header


def test_analyze_polish_returns_polished_doc_as_download(monkeypatch) -> None:
    def _fake_agent(chunk: str):
        return {"revised_text": chunk + "（润色）", "notes": []}

    monkeypatch.setattr(pipeline, "call_agent_and_normalize", _fake_agent)

    result = analyze(mode="text", text="关于测试通知\n请按时完成。", do_polish=True)
    assert result["download_filename"] == "润色后文档.docx"
    assert "润色后文档.docx" in result["files"]
    assert result["download_bytes"] == result["files"]["润色后文档.docx"]


def test_load_agent_config_supports_assignment_style(tmp_path: Path, monkeypatch) -> None:
    from agent_config import load_agent_config

    cfg_file = tmp_path / "agent_config.local.json"
    cfg_file.write_text(
        'AUTH_KEY = "demo_key".strip()\nAUTH_SECRET = "demo_secret".strip()\nBASE_URL = "https://uat.agentspro.cn"\nAGENT_ID = "demo_agent"\n',
        encoding="utf-8",
    )

    monkeypatch.setenv("AGENT_CONFIG_PATH", str(cfg_file))
    cfg = load_agent_config()

    assert cfg.auth_key == "demo_key"
    assert cfg.auth_secret == "demo_secret"
    assert cfg.base_url == "https://uat.agentspro.cn"
    assert cfg.agent_id == "demo_agent"


def test_load_agent_config_json_and_env_override(tmp_path: Path, monkeypatch) -> None:
    from agent_config import load_agent_config

    cfg_file = tmp_path / "agent_config.local.json"
    cfg_file.write_text('{"AUTH_KEY": "file_key", "AUTH_SECRET": "file_secret", "BASE_URL": "https://file-host", "AGENT_ID": "file_agent"}', encoding="utf-8")

    monkeypatch.setenv("AGENT_CONFIG_PATH", str(cfg_file))
    monkeypatch.setenv("AUTH_SECRET", "env_secret")
    monkeypatch.setenv("AGENT_BASE_URL", "https://env-host")

    cfg = load_agent_config()

    assert cfg.auth_key == "file_key"
    assert cfg.auth_secret == "env_secret"
    assert cfg.base_url == "https://env-host"
    assert cfg.agent_id == "file_agent"


def test_load_agent_config_supports_lowercase_keys(tmp_path: Path, monkeypatch) -> None:
    from agent_config import load_agent_config

    cfg_file = tmp_path / "agent_config.local.json"
    cfg_file.write_text('{"auth_key": "k", "auth_secret": "s", "base_url": "https://lower", "agent_id": "a"}', encoding="utf-8")

    monkeypatch.setenv("AGENT_CONFIG_PATH", str(cfg_file))
    cfg = load_agent_config()

    assert cfg.auth_key == "k"
    assert cfg.auth_secret == "s"
    assert cfg.base_url == "https://lower"
    assert cfg.agent_id == "a"


def test_has_agent_failure_detects_agent_notes() -> None:
    from main import _has_agent_failure

    assert _has_agent_failure([{"chunk_index": 0, "notes": [{"type": "agent", "message": "大模型调用失败：xxx"}]}])
    assert not _has_agent_failure([{"chunk_index": 0, "notes": [{"type": "agent", "message": "未配置智能体鉴权参数"}]}])
    assert not _has_agent_failure([{"chunk_index": 0, "notes": []}])


def test_has_agent_failure_only_flags_real_failures() -> None:
    from main import _has_agent_failure

    assert _has_agent_failure([{"chunk_index": 0, "notes": [{"type": "agent", "message": "stream_error"}]}])
    assert _has_agent_failure([{"chunk_index": 0, "notes": [{"type": "agent", "message": "http_error"}]}])
    assert not _has_agent_failure([{"chunk_index": 0, "notes": [{"type": "agent", "message": "未配置智能体鉴权参数，已跳过大模型润色。"}]}])
