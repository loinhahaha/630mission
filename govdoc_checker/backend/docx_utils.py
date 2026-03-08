from __future__ import annotations

import os
import subprocess
from typing import List, Any

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

import shutil
import sys


def _find_soffice() -> str:
    for name in ("soffice", "soffice.exe"):
        p = shutil.which(name)
        if p:
            return p

    if sys.platform.startswith("win"):
        candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for c in candidates:
            if os.path.exists(c):
                return c
    return "soffice"


def _run(cmd: list[str]) -> tuple[int, str, str]:
    p = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    return p.returncode, p.stdout, p.stderr


def load_input_as_docx(input_path: str, workdir: str) -> str:
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".docx":
        return input_path
    if ext == ".doc":
        rc, _out, err = _run([_find_soffice(), "--headless", "--convert-to", "docx", "--outdir", workdir, input_path])
        base = os.path.splitext(os.path.basename(input_path))[0]
        candidate = os.path.join(workdir, base + ".docx")
        if rc == 0 and os.path.exists(candidate):
            return candidate
        raise RuntimeError(f"无法将 .doc 转为 .docx。请安装 LibreOffice。stderr={err}")
    raise RuntimeError(f"不支持的文件类型: {ext}（仅支持 .docx/.doc）")


def extract_docx_text(docx_path: str) -> List[str]:
    # 读取 Word 并按段落抽取文本。
    # 注意：这里刻意保留空段，否则规则输出的 paragraph_index 会与原文错位。
    doc = Document(docx_path)
    # 保留空段，确保 paragraph_index 与原文段落严格对齐
    return [(p.text or "").rstrip() for p in doc.paragraphs]


def _apply_gb9704_page_setup(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(37)
    sec.left_margin = Mm(28)
    sec.right_margin = Mm(26)
    sec.bottom_margin = Mm(35)


def _set_doc_default_font(doc: Document, font_name: str = "仿宋", font_size_pt: int = 16) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(font_size_pt)


def build_final_docx_from_text(text: str, out_path: str, as_input: bool = False) -> str:
    doc = Document()
    _apply_gb9704_page_setup(doc)
    _set_doc_default_font(doc, "方正仿宋_GBK", 16)

    lines = [ln.rstrip() for ln in (text or "").splitlines()]
    lines = [ln for ln in lines if ln.strip() != ""]
    if lines and len(lines[0]) <= 30 and "：" not in lines[0]:
        doc.add_paragraph("")
        title = lines.pop(0).strip()
        p = doc.add_paragraph(title)
        p.alignment = 1
        run = p.runs[0]
        run.font.name = "方正小标宋_GBK"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "方正小标宋_GBK")
        run.font.size = Pt(22)
        doc.add_paragraph("")

    for ln in lines:
        p = doc.add_paragraph(ln.strip())
        p.paragraph_format.first_line_indent = Pt(32)
        p.paragraph_format.line_spacing = Pt(29)

    doc.save(out_path)
    return out_path



def _apply_highlight_with_insertions(paragraph, text: str, highlights: list[bool], insertions_after: dict[int, list[str]]) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    if not text:
        return

    i = 0
    while i < len(text):
        j = i + 1
        while j < len(text) and highlights[j] == highlights[i]:
            j += 1
        run = paragraph.add_run(text[i:j])
        if highlights[i]:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        if j in insertions_after:
            for msg in insertions_after[j]:
                note_run = paragraph.add_run(msg)
                note_run.font.size = Pt(10)
                note_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        i = j



def _find_span_from(text: str, snippet: str, start_pos: int = 0) -> tuple[int, int] | None:
    snippet = (snippet or "").strip()
    if not snippet:
        return None
    i = text.find(snippet, max(0, start_pos))
    if i < 0:
        return None
    return i, i + len(snippet)


def _issue_note(issue: Any) -> str:
    msg = (getattr(issue, "message", "") or "").strip()
    if not msg:
        msg = (getattr(issue, "issue_type", "问题") or "问题").strip()
    return f"【{msg}】"


def build_annotated_original_docx(original_docx_path: str, issues: list[Any], out_path: str) -> str:
    # 以“原文不改结构”为原则进行标注：
    # 1) 按 paragraph_index 聚合问题
    # 2) 按 evidence 尽量精确定位错误片段
    # 3) 错误处高亮 + 批注，不再拼接到段落正文里
    doc = Document(original_docx_path)
    by_para: dict[int, list[Any]] = {}
    for it in issues:
        idx = getattr(it, "paragraph_index", None)
        if idx is None:
            continue
        by_para.setdefault(idx, []).append(it)

    def mark_paragraph(paragraph, para_issues: list[Any]) -> None:
        text = paragraph.text or ""
        if not text:
            return

        highlights = [False] * len(text)
        insertions_after: dict[int, list[str]] = {}
        cursor = 0

        for it in para_issues:
            ev = getattr(it, "evidence", None)
            span = _find_span_from(text, ev, cursor)

            if span is None and getattr(it, "issue_type", "") == "空格":
                for k, ch in enumerate(text):
                    if ch in (" ", "　", "	"):
                        span = (k, k + 1)
                        break

            if span is None and text:
                span = (len(text) - 1, len(text))

            if span is not None:
                for i in range(span[0], span[1]):
                    highlights[i] = True
                cursor = span[1]
                insertions_after.setdefault(span[1], []).append(_issue_note(it))

        _apply_highlight_with_insertions(paragraph, text, highlights, insertions_after)

    for i, p in enumerate(list(doc.paragraphs)):
        if i in by_para:
            mark_paragraph(p, by_para[i])

    doc.save(out_path)
    return out_path


def try_convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    outdir = os.path.dirname(pdf_path)
    base = os.path.splitext(os.path.basename(docx_path))[0]
    try:
        rc, _out, _err = _run([_find_soffice(), "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path])
        candidate = os.path.join(outdir, base + ".pdf")
        if rc == 0 and os.path.exists(candidate):
            if candidate != pdf_path:
                os.replace(candidate, pdf_path)
            return True
    except FileNotFoundError:
        return False
    except Exception:
        return False
    return False
